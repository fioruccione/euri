#!/usr/bin/env python3
"""Regressioni pure per artefatti di sessione e documenti verificati."""
from __future__ import annotations

import json
import hashlib
import tempfile
import threading
from pathlib import Path
from types import SimpleNamespace

from agent.executor import Executor, ToolResult
from agent.tools import document_composer, text_writer
from core.action_controller import (
    ActionAuthority, ActionDecision, ActionDisposition, ActionProposal,
)


PLAN = {
    "title": "Relazione di prova",
    "subtitle": "Versione revisionata",
    "blocks": [
        {"type": "heading", "level": 1, "text": "Sintesi"},
        {"type": "paragraph", "text": "Il dato verificato è 42 °C."},
        {"type": "bullet_list", "items": ["Primo punto", "Secondo punto"]},
        {
            "type": "table",
            "headers": ["Parametro", "Valore"],
            "rows": [["Temperatura", "42 °C"]],
        },
    ],
}


class _JsonChat:
    def __init__(self):
        self.calls = []

    def chat(self, **kwargs):
        self.calls.append(kwargs)
        return SimpleNamespace(
            message=SimpleNamespace(content=json.dumps(PLAN, ensure_ascii=False))
        )


class _Brain:
    def __init__(self):
        self.history_lock = threading.Lock()
        self._conversation_history = [
            {"role": "assistant", "content": "Renderei più chiara la sintesi."}
        ]
        self.injected = []

    def inject_tool_result(self, *args):
        self.injected.append(args)


class _DirectController:
    def __init__(self, proposal):
        self.proposal = proposal

    def propose(self, *_args, **_kwargs):
        return self.proposal

    def decide(self, proposal, _capabilities):
        return ActionDecision(ActionDisposition.EXECUTE, proposal, "test")


class _AbstainController:
    def propose(self, *_args, **_kwargs):
        return None

    def decide(self, proposal, _capabilities):
        return ActionDecision(ActionDisposition.ABSTAIN, proposal, "no_proposal")


def test_plan_is_structured_and_grounded_in_full_source():
    chat = _JsonChat()
    source = "Bozza completa con il dato 42 °C e il nome Gio Style."
    plan = document_composer.build_document_plan(
        source,
        "Riorganizza senza aggiungere fatti.",
        recent_context="ASSISTANT: usa sezioni più chiare",
        chat=chat,
        model="fake",
    )
    assert plan == PLAN
    prompt = chat.calls[0]["messages"][0]["content"]
    assert source in prompt
    assert "usa sezioni più chiare" in prompt
    assert chat.calls[0]["format"] == "json"


def test_txt_docx_pdf_are_reopened_and_verified():
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        for fmt in ("txt", "docx", "pdf"):
            receipt = document_composer.render_document(
                PLAN, root, fmt=fmt, filename="relazione"
            )
            path = Path(receipt["filepath"])
            assert path.exists() and path.stat().st_size == receipt["bytes"]
            assert receipt["format"] == fmt
            assert len(receipt["sha256"]) == 64
            assert receipt["validation"]["nonempty"] is True
        assert (root / "relazione.txt").read_text(encoding="utf-8").count("42 °C") == 2


def test_existing_output_is_never_overwritten():
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        first = document_composer.render_document(PLAN, root, fmt="txt", filename="bozza")
        second = document_composer.render_document(PLAN, root, fmt="txt", filename="bozza")
        assert first["filepath"] != second["filepath"]
        assert Path(first["filepath"]).exists() and Path(second["filepath"]).exists()


def test_clipboard_preview_keeps_complete_operational_artifact():
    original = text_writer._read_clipboard_text
    full = "A" * 3792
    try:
        text_writer._read_clipboard_text = lambda: full
        result = text_writer.tool_clipboard_read({})
    finally:
        text_writer._read_clipboard_text = original
    assert result.success is True
    assert len(result.output) < 400
    assert result.raw_data["artifact_content"] == full
    executor = Executor.__new__(Executor)
    executor._session_artifact_lock = threading.Lock()
    executor._session_artifact = None
    executor._capture_session_artifact("clipboard_read", result)
    assert executor.get_session_artifact()["content"] == full
    executor._capture_session_artifact(
        "clipboard_read", ToolResult(False, "appunti non disponibili")
    )
    assert executor.get_session_artifact() is None


def test_semantic_dispatch_uses_exact_request_and_returns_real_receipt():
    with tempfile.TemporaryDirectory() as tmp:
        executor = Executor()
        executor.brain = _Brain()
        executor._session_artifact = {
            "id": "artifact:test",
            "kind": "uploaded_documents",
            "source": "ui",
            "filenames": ["bozza.docx"],
            "content": "Contenuto completo della bozza: 42 °C.",
            "captured_at": __import__("time").time(),
        }
        received = {}
        original_builder = document_composer.build_document_plan
        original_handler = executor._registry["compose_document"].handler
        try:
            document_composer.build_document_plan = lambda source, instruction, **kwargs: (
                received.update({"source": source, "instruction": instruction}) or PLAN
            )
            executor._registry["compose_document"].handler = lambda params, **kwargs: (
                document_composer.compose_document_tool(
                    params,
                    artifact=executor.get_session_artifact(),
                    recent_context=executor._recent_document_context(),
                    output_dir=Path(tmp),
                )
            )
            proposal = ActionProposal(
                capability="executor.compose_document",
                args={"instruction": "parafrasi non autorizzata", "format": "docx"},
                target_id=None,
                authority=ActionAuthority.USER_EXPLICIT,
                confidence=0.99,
            )
            utterance = "Genera direttamente un documento Word con le modifiche suggerite."
            result = executor.dispatch_contextual_action(
                utterance,
                previous_euri_turn="Ho suggerito una sintesi più chiara.",
                controller=_DirectController(proposal),
            )
        finally:
            document_composer.build_document_plan = original_builder
            executor._registry["compose_document"].handler = original_handler
        assert result["success"] is True
        assert received["instruction"] == utterance
        assert "42 °C" in received["source"]
        assert Path(result["raw_data"]["filepath"]).exists()
        assert result["raw_data"]["validation"]["nonempty"] is True
        assert executor.brain.injected


def test_semantic_abstain_is_fail_closed_not_chat():
    executor = Executor()
    result = executor.dispatch_contextual_action(
        "Preparamelo e salvalo.", controller=_AbstainController()
    )
    assert result["success"] is False
    assert result["fail_closed"] is True
    assert "Non ho eseguito nulla" in result["output"]


def test_conservative_docx_revision_preserves_layout_footer_and_lists():
    from docx import Document
    from docx.shared import Cm

    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        source = root / "originale.docx"
        doc = Document()
        section = doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(1.9)
        section.footer.paragraphs[0].text = "Lucy Plast — footer legale"
        doc.add_heading("Conformità", level=1)
        doc.add_paragraph("Testo originale da aggiornare.")
        doc.add_paragraph("Voce elenco intatta", style="List Bullet")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Parametro"
        table.cell(0, 1).text = "Valore"
        table.cell(1, 0).text = "Temperatura"
        table.cell(1, 1).text = "40 °C"
        doc.save(source)
        source_sha = hashlib.sha256(source.read_bytes()).hexdigest()
        receipt = document_composer.revise_docx(
            source,
            {
                "summary": "Aggiornato il solo passaggio autorizzato",
                "warnings": ["Misura formulata come proposta"],
                "edits": [{
                    "paragraph_index": 1,
                    "replacement": "Testo revisionato come proposta, non come fatto operativo.",
                    "reason": "richiesta utente",
                }],
                "table_edits": [{
                    "table_index": 0,
                    "row": 1,
                    "column": 1,
                    "replacement": "42 °C",
                    "reason": "dato autorizzato",
                }],
            },
            root / "output",
            expected_sha256=source_sha,
        )
        revised = Document(receipt["filepath"])
        assert revised.sections[0].page_width == section.page_width
        assert revised.sections[0].page_height == section.page_height
        assert revised.sections[0].left_margin == section.left_margin
        assert revised.sections[0].footer.paragraphs[0].text == "Lucy Plast — footer legale"
        assert revised.paragraphs[2].style.name == "List Bullet"
        assert revised.paragraphs[1].text.startswith("Testo revisionato")
        assert revised.tables[0].cell(1, 1).text == "42 °C"
        assert receipt["validation"]["structure_preserved"] is True
        assert Document(source).paragraphs[1].text == "Testo originale da aggiornare."


def test_conservative_docx_revision_rejects_stale_source():
    from docx import Document

    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        source = root / "originale.docx"
        doc = Document()
        doc.add_paragraph("Originale")
        doc.save(source)
        try:
            document_composer.revise_docx(
                source,
                {"edits": [{"paragraph_index": 1, "replacement": "Nuovo"}]},
                root / "output",
                expected_sha256="0" * 64,
            )
        except ValueError as exc:
            assert "cambiato dopo la lettura" in str(exc)
        else:
            raise AssertionError("stale source non bloccata")


def test_compose_tool_uses_conservative_path_for_real_docx():
    from docx import Document

    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        source = root / "lettera.docx"
        doc = Document()
        doc.add_paragraph("Testo iniziale")
        doc.save(source)
        original_builder = document_composer.build_docx_edit_plan
        try:
            document_composer.build_docx_edit_plan = lambda *args, **kwargs: {
                "summary": "Una modifica puntuale",
                "warnings": [],
                "edits": [{
                    "paragraph_index": 0,
                    "replacement": "Testo revisionato",
                    "reason": "test",
                }],
                "table_edits": [],
            }
            result = document_composer.compose_document_tool(
                {"instruction": "Aggiorna il testo", "format": "docx"},
                artifact={
                    "id": "artifact:docx",
                    "filename": source.name,
                    "source_path": str(source),
                    "content": "Testo iniziale",
                    "sha256": hashlib.sha256(source.read_bytes()).hexdigest(),
                    "version": 3,
                    "kind": "uploaded_document",
                },
                output_dir=root / "output",
            )
        finally:
            document_composer.build_docx_edit_plan = original_builder
        assert result.success is True
        assert result.raw_data["mode"] == "conservative_revision"
        assert result.raw_data["source_version"] == 3
        assert "revisionato e verificato" in result.output
        assert Document(result.raw_data["filepath"]).paragraphs[0].text == "Testo revisionato"
        assert Document(source).paragraphs[0].text == "Testo iniziale"


if __name__ == "__main__":
    test_plan_is_structured_and_grounded_in_full_source()
    test_txt_docx_pdf_are_reopened_and_verified()
    test_existing_output_is_never_overwritten()
    test_clipboard_preview_keeps_complete_operational_artifact()
    test_semantic_dispatch_uses_exact_request_and_returns_real_receipt()
    test_semantic_abstain_is_fail_closed_not_chat()
    test_conservative_docx_revision_preserves_layout_footer_and_lists()
    test_conservative_docx_revision_rejects_stale_source()
    test_compose_tool_uses_conservative_path_for_real_docx()
    print("test_document_composer: OK")
